from docx import Document
import os
import pandas
import re
from tqdm import tqdm
import sys

def check_underline(par):
    try:
        text = ''
        for run in par.runs:
            if run.underline:
                text += par.text
        if text.strip() == par.text.strip():
            return True
        
        # it could be that the par run.underline == None 
        #witch means it has inherted the value from the style
        current_style = par.style
        while current_style:
            if current_style.font.underline:
                return True

            # Move up the style hierarchy
            current_style = current_style.base_style

        return False
    except Exception as e:
        print (f'Exception in check_underline: {e}')

def get_all_docx_in_current_foleder(folder_path):
    try:

        info =[]
        # go throw all the documants in the directory
        current_path =os.path.join(os.getcwd(),folder_path)
        
        for file_number,filename in enumerate(os.listdir(current_path)):
            if filename.endswith('docx'):
                
                attributes = filename.split('_')# get the attributes of the file
                #then assigen it to the right variable
                info.append({})
                info[-1]['file_name'] = filename
                info[-1]['number'] = int(attributes[0])# the XX number 

                if attributes[1] == 'ptv':#committee or plenary
                    info[-1]['type'] = 'committee'
                elif (attributes[1] == 'ptm'):
                    info[-1]['type'] = 'plenary'

                text = Document(os.path.join(current_path,filename))
                #clear the text from B,
                '''for par_number, par in enumerate (text.paragraphs):
                    if par.text != '' and par.text[0] == '<': #some the senteces begin with "<" and ends with ">" and does not appear in the document so we remvove them 
                        par.text = par.text[1:-1]
                        text.paragraphs[par_number] = par'''
                info[-1]['text'] = text
                info[-1]['debugging_number'] = attributes[2]
        return info
    except Exception as e:
        print(f'Exception in get_all_docx_in_current_foleder: {e}')

def clear_name(name): #input a string (name) output the cleared name
    try:
        name = name.strip()
        componints = name.split(' ')
        new_name = ""
        open_parentheses =False 
        for comp in componints:
            if comp == '':
                continue

            if '(' in comp:
                open_parentheses = False
                continue
            if open_parentheses:
                continue
            if '"' not in comp and '”' not in comp:#this means that the component is not a short cut
                if '\'' == comp[-1] and len(comp) <4: #then this means that it might be the person numbering of his possition
                    new_name = '' # is so then just throw all of what was before it
                    continue
                    #the problem here if the number is more than 2 digits in hebrew this means this code will not capture it
                    #another problem if someone name has 2 latters and ends with '\'' 
                    #another problem is if his posstion came after his name
                if ")" in comp: #if we have open_parentheses then we should not inclut what is in it
                    if "(" in comp:# if we have closing parentheses then we just dont take this comp
                        continue
                    else:# else we must take the following comps until we see closing parentheses
                        open_parentheses = True
                elif comp == "-" or comp == '–' or comp == '~' or comp ==',':#if the name has " - " then take the first part
                    break
                else:
                    new_name += comp+" "
        new_name = new_name.strip()
        if ',' in new_name:
            return ''
        if new_name != "" and new_name.find(':')+1 == len(new_name):
            new_name = new_name[:-1]

        return new_name.strip()
    except Exception as e:
        print(f'Exception in clear_name: {e}')



def clean_text(text):
    try:
        if text == '':
            return ''
        # cleaning the text of the charachters that cant be in it
        pattern_with_space = re.compile('[א-ת0-9!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~ ]+')
        matches = pattern_with_space.findall(text)
        cleaned_text = ''.join(matches)
        #if after cleaning the none hebrew words and it is not equvlent then it must be that we have non hebrew text witch means we nust delete
        if cleaned_text != text:
            return ''
        #now we have cleared all the charetars that can be like English letters ,Japanese .. 

        '''non_alphabetic_characters = ['!', '"', '#', '$', '%', '&', '\'', '(', ')', '*',
                                '+', ',', '-', '.', '/', ':', ';', '<', '=', '>',
                                '?', '@', '[', '\\', ']', '^', '_', '`', '{', '|',
                                '}', '~'] '''
        hebrew_letters = [chr(code) for code in range(0x05D0, 0x05EA + 1)]

        hebrew = False
        for latter in hebrew_letters:
            if latter in cleaned_text:
                hebrew=True
                break
        
        if hebrew == False: #if we dont have any hebrew latter then just delete the sentence
            return ''
        

        if '- - -'  in cleaned_text:
            return ''
        '''
        new_text = ''
        i = 0 
        flag_hebrew = False # if we dont have any hebrew then this is not a text
        while i< len(cleaned_text):
            if flag_hebrew == False and cleaned_text[i] in hebrew_letters:
                flag_hebrew = True
            elif flag_hebrew == False:
                i+=1
            if flag_hebrew == True: # the first charecter must be hebrew
                delelte_last = False
                
                while i!=len(cleaned_text) and cleaned_text[i] in non_alphabetic_characters and get_the_next_index_not_space(cleaned_text,i) != len(cleaned_text) and cleaned_text[get_the_next_index_not_space(cleaned_text,i)] in non_alphabetic_characters:
                    i=get_the_next_index_not_space(cleaned_text,i)
                    delelte_last = True
                if delelte_last:
                    i=get_the_next_index_not_space(cleaned_text,i)
                if i!= len(cleaned_text):
                    new_text+=cleaned_text[i]
                    i+=1
        '''
        return text
    except Exception as e:
        print(f'Exception in clean_text {e}')

def get_the_next_index_not_space(text,i):
    for j in range(i+1,len(text)):
        if text[j] != ' ':
            return j
    return len(text)
    
def split_paragrph(par):
    try:
        new_text = ''
        sentence_seprator = '.؟!:;' #we split according to thoughs
        we_are_in_quets = False
        par_parts = par.text.strip().split(' ')
        sentece_list = []
        for part in par_parts:
            
            if part =='':# if the text is empty there is nothing to do 
                continue
            new_text += part +" "
            if '"' == part[0] and we_are_in_quets == False: 
                we_are_in_quets = True
            if '"' == part[-1] or (len(part)>=2 and part[-2] == '"' and part[-1] in sentence_seprator):
                # if the last charerchter is a quation mark or the second to last charechter and it has a seprator then this means that this is the end of the quation
                we_are_in_quets = False

            if part[-1]  in sentence_seprator or (len(part)>=2 and part[-2] in sentence_seprator):
                # if the sentence ends with are the sentence_seprator and then '\"' we catch that 
                if we_are_in_quets == False:
                    # we just add if the sentence_seprator is not in the quets
                    sentece_list.append(new_text.strip())
                    new_text = ''
        #if we start a quet but didnt end it then save the text
        if we_are_in_quets:
            sentece_list.append(new_text)
        return sentece_list
    except Exception as e:
        print(f'exception in split_paragrph: {e}')

                

'''
def split_paragrph(par):
    pattern = re.compile(r'[.؟!:;]')
    
    sentences =  pattern.split(par.text)

    results = []
    #delelte epmty strings
    for sentence in sentences:
        if sentence.strip() == '':
            continue
        elif len(sentence.strip().split())>=4: #may change
            results.append(sentence.strip())
    return results
'''

def make_token(list_text):
    try:
        token_list= []
        non_alphabetic_characters = ['!', '"', '#', '$', '%', '&', '\'', '(', ')', '*',
                                '+', ',', '-', '.', '/', ':', ';', '<', '=', '>',
                                '?', '@', '[', '\\', ']', '_', '`', '{', '|',
                                '}', '~'] # all the non_alphabetic_characters that may the writer use
        for text in list_text:
            current_token_list =[]
            text_parts = text.split(' ')
            for j in range(len(text_parts)):
                # for each word if the end is a 
                word = text_parts[j]
                for i in range(len(word)):
                    #if we have non_alphabetic_characters in the start until hebrew charechter make a new token
                    if word[i] in non_alphabetic_characters:
                        current_token_list.append(word[i])
                        text_parts[j] = text_parts[j][1:] # delete from the orginal text
                    else:
                        break
                for i in range(len(word)):
                    # the same as the past loop but from the end back to an hebrew latter
                    if word[len(word)-i-1] in non_alphabetic_characters:
                        current_token_list.append(len(word)-i-1)
                        text_parts[j] = text_parts[j][:-1]
                    else:
                        break
            if len(current_token_list) + len (text_parts)<4: #if we have less than 4 tokens doelete
                continue
            # add to the data
            token_list.extend(current_token_list)
            token_list.extend(text_parts)

        return token_list
    except Exception as e:
        print(f'Exception in make_token: {e}')

#if len(sys.argv)>2:
#    sys.exit(1)

if __name__ =="__main__":
    try:
        # taking the arguments
        if len(sys.argv) !=3:
            print('must be 2 args')
            sys.exit(1)
        folder_path = sys.argv[1]
        return_path = sys.argv[2]
        data = get_all_docx_in_current_foleder(folder_path)
        temp_data = data
        for docx_number, docx in tqdm(enumerate(temp_data), total=len(temp_data)):
            if docx['type'] == 'plenary':
                speaker_name ='' #the name of the current_speaker
                speaker_text = {}# a dictionary of all the availabe text (speakre name,all speaker's text)
                for par in docx['text'].paragraphs:
                    symbol_index = par.text.strip().find(":")
                    #if we arrive at a sentence that has ':' in the end and all of it underlined then this is a speaker
                    if symbol_index>=0 and symbol_index== len(par.text.strip()) -1 and check_underline(par):
                        new_speaker_name = clear_name(par.text.strip())
                        if new_speaker_name == '':# if there is ',' in the sentnce then it is not a name so we gove ot to the past speaker
                            # we split the paragraph to senteces
                            splited_text = split_paragrph(par)
                            text =[] 
                            #clean each sentence and if the sentece is removes dont append
                            for text_part in splited_text:
                                cleaned_texted_part = clean_text(text_part)
                                if cleaned_texted_part != '':
                                    text.append(cleaned_texted_part)
                        # add the text to the speaker after making sure it is not empty and clean
                            if speaker_name != '':
                            
                                speaker_text[speaker_name].extend(make_token(text))
                        else:
                            speaker_name = new_speaker_name #if there is no problem with new name make it the current speaker
                        if speaker_name not in list(speaker_text.keys()):
                            #ad the speaker to the dictionary if not yet added
                            speaker_text[speaker_name] = []

                    elif speaker_name!='':# if we have a speaker then add the text to his name
                        splited_text = split_paragrph(par)
                        # we do it like before split ,clean ,add 
                        text =[] 
                        for text_part in splited_text:
                            cleaned_texted_part = clean_text(text_part)
                            if cleaned_texted_part != '':
                                text.append(cleaned_texted_part)        
                        speaker_text[speaker_name].extend(make_token(text))

                data[docx_number]['speaker_data'] = speaker_text# save the data

            if docx['type'] == 'committee':
                for par_number, par in enumerate (docx['text'].paragraphs):
                    if par.text != '' and par.text[0] == '<': #some the senteces begin with "<" and ends with ">" and does not appear in the document so we remvove them 
                        par.text = par.text[1:-1]
                first_subject = ''#the first subject the document is talking about
                take_names = 0#a flag to list anither name or not
                i = -1
                for par in docx['text'].paragraphs:
                    i+=1
                    #find the first_subject this helps us to get all the people how are in the documant
                    if first_subject =='' and  par.text.strip().find('סדר') in [0,1] and par.text.find('יום')>=0 :

                        if len(par.text)>=12: #if this happens then we have the first subject in the same line as the current par
                            if par.text.find(':יום')>=0:# if it has ":" then split with : other wise split with " "
                                first_subject = par.text.split(":", 2)[1]  
                            else:
                                first_subject = par.text.split(" ", 2)[2]
                            first_subject = first_subject.split(',')[0]#if the sentence has more than one subject take the first
                        else:# the subject not in the paragraph
                            current_i = i+1
                            while docx['text'].paragraphs[current_i].text.strip()== '':
                                current_i+=1
                            first_subject = docx['text'].paragraphs[current_i].text
                            first_subject = first_subject.split(',')[0].strip()#if the sentence has more than one subject take the first 
                    
                        

                speaker_name ='' # like in the past section
                speaker_text = {}
                first_subject_counter = 2
                for par_number, par in enumerate (docx['text'].paragraphs):
                    # some times the chairman opens before the fire subject apear
                    if first_subject_counter>0 and 'יו"ר' in par.text and ":" in par.text:
                        first_subject_counter = 0
                    if first_subject_counter>0 and len(set(par.text.split(' ')).intersection(first_subject.split(' ')))>=(3*len(first_subject.split(' ')))/4:# if we have more than 3/4 of the words interact this is probably the first subject 
                        first_subject_counter -=1
                    if first_subject_counter ==0:
                        #begin to read each speaker data
                        symbol_index = par.text.strip().find(":")
                        #like the previous kind if we find : and unerline then we say that this is a name
                        if symbol_index>=0 and symbol_index== len(par.text.strip()) -1 and check_underline(par):
                            new_speaker_name = clear_name(par.text.strip())
                            if new_speaker_name == '':
                                # the function return empty sting if the sentece is not a name
                                # and then like the previous type we split clean and append
                                splited_text = split_paragrph(par)
                                text =[] 
                                for text_part in splited_text:
                                    cleaned_texted_part = clean_text(text_part)
                                    if cleaned_texted_part != '':
                                        text.append(cleaned_texted_part)
                            # add the text to the speaker after making sure that speaker is not empty is not empty and clean
                                if speaker_name != '':
                                    
                                    speaker_text[speaker_name].extend(make_token(text))
                            else:
                                speaker_name = new_speaker_name #like the previous kind if the name is a real name then it is the speaker
                            # add it to the dictionary if he doesnt have an entry
                            if speaker_name not in list(speaker_text.keys()):
                                speaker_text[speaker_name] = []
                        elif speaker_name!='':
                            # we just add data if we have a speaker
                            #and then like before split, clean and append each sentence
                            splited_text = split_paragrph(par)
                            text =[] 
                            for text_part in splited_text:
                                cleaned_texted_part = clean_text(text_part)
                                if cleaned_texted_part != '':
                                    text.append(cleaned_texted_part)        
                            speaker_text[speaker_name].extend(make_token(text))
                #sav e the data of the current file
                data[docx_number]['speaker_data'] = speaker_text

                            
                
                #end of new code
                #print('documant number = '+docx['debugging_number'])
                #print(first_subject[::-1])


        #creating the csv
        columns_name = {'protocol_name':[],'knesset_number':[],'protocol_type':[],'speaker_name':[],'sentence_text':[]}
        list_of_speakers = []
        #intilize the data frame
        df = pandas.DataFrame(columns_name)
        for row_index, row in tqdm(enumerate(data), total=len(data)):
            #make a new row for each file
            new_row = {'protocol_name':row['file_name'],'knesset_number':row['number'],'protocol_type':row['type']}
            for speaker in row['speaker_data'].keys():
                #fill the missing data for each speaker
                new_row['speaker_name'] = speaker
                new_row ['sentence_text'] = row['speaker_data'][speaker]
                new_df = pandas.DataFrame(new_row)
                new_df=new_df.ffill()
                #save the data in the overall data frame
                df = pandas.concat([df, new_df])
            ''' list_of_speakers.append(speaker)
                for text in row['speaker_data'][speaker]:
                    new_row ['speaker_name'] = speaker
                    new_row ['sentence_text'] = text
                    df.loc[len(df)] = new_row'''
            

        #convert

        df.to_csv(return_path,index = False ,encoding='utf-8')
        #df.to_excel('our_data.xlsx',index = False)    

        print('csv = ' + str(len(df)))

    except Exception as e:
        # Handle any exception
        print(f"An error occurred in main: {e}")